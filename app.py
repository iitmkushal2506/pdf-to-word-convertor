from flask import Flask, render_template, request, send_file, jsonify
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import uuid
import tempfile

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

UPLOAD_FOLDER = tempfile.gettempdir()


# ── Conversion helpers (same logic as standalone script) ──────────────────────

def is_heading(text, font_size, avg_size):
    if not text or len(text.strip()) < 2:
        return False, 0
    stripped = text.strip()
    if font_size and avg_size:
        ratio = font_size / avg_size
        if ratio >= 1.5 and len(stripped) < 120:
            return True, 1
        if ratio >= 1.2 and len(stripped) < 120:
            return True, 2
    if stripped.isupper() and len(stripped) < 80 and not stripped.endswith('.'):
        return True, 2
    return False, 0


def get_avg_font_size(page):
    sizes = [c["size"] for c in page.chars if c.get("size")]
    return sum(sizes) / len(sizes) if sizes else 12.0


def detect_alignment(line_chars, page_width):
    if not line_chars:
        return "left"
    x_vals = [c["x0"] for c in line_chars]
    avg_x = sum(x_vals) / len(x_vals)
    center = page_width / 2
    if abs(avg_x - center) < page_width * 0.15:
        return "center"
    if avg_x > page_width * 0.6:
        return "right"
    return "left"


def add_table_to_doc(doc, table_data):
    if not table_data:
        return
    num_cols = max(len(row) for row in table_data)
    word_table = doc.add_table(rows=len(table_data), cols=num_cols)
    word_table.style = "Table Grid"
    for r_idx, row in enumerate(table_data):
        for c_idx, cell_text in enumerate(row):
            cell = word_table.cell(r_idx, c_idx)
            cell.text = cell_text or ""
            if r_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()


def convert_pdf_to_word(pdf_path, output_path, detect_tables=True):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            table_bboxes = []
            if detect_tables:
                tables = page.extract_tables()
                for tbl in tables:
                    if tbl:
                        add_table_to_doc(doc, tbl)
                for tbl_obj in page.find_tables():
                    table_bboxes.append(tbl_obj.bbox)

            avg_size = get_avg_font_size(page)
            words = page.extract_words(extra_attrs=["size", "fontname"], keep_blank_chars=False)

            def in_table(word):
                for (x0, top, x1, bottom) in table_bboxes:
                    if word["x0"] >= x0 and word["x1"] <= x1 and \
                       word["top"] >= top and word["bottom"] <= bottom:
                        return True
                return False

            filtered_words = [w for w in words if not in_table(w)]
            lines = {}
            for w in filtered_words:
                key = round(w["top"], 1)
                lines.setdefault(key, []).append(w)

            prev_bottom = 0.0
            for top_pos in sorted(lines.keys()):
                line_words = sorted(lines[top_pos], key=lambda w: w["x0"])
                line_text  = " ".join(w["text"] for w in line_words).strip()
                if not line_text:
                    continue
                if prev_bottom and (top_pos - prev_bottom) > avg_size * 1.8:
                    doc.add_paragraph()
                font_size = line_words[0].get("size") if line_words else None
                heading_flag, level = is_heading(line_text, font_size, avg_size)
                if heading_flag:
                    heading_map = {1: "Heading 1", 2: "Heading 2"}
                    doc.add_paragraph(line_text, style=heading_map.get(level, "Heading 2"))
                else:
                    para = doc.add_paragraph(line_text)
                    if font_size:
                        for run in para.runs:
                            run.font.size = Pt(round(font_size))
                    align = detect_alignment(
                        [c for c in page.chars
                         if round(c.get("top", -1), 1) == top_pos and not in_table(c)],
                        page.width,
                    )
                    if align == "center":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align == "right":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                prev_bottom = line_words[-1].get("bottom", top_pos) if line_words else top_pos

            if page_num < total_pages:
                doc.add_page_break()

    doc.save(output_path)
    return total_pages


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/convert', methods=['POST'])
def convert():
    if 'pdf' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['pdf']
    if not file.filename.endswith('.pdf'):
        return jsonify({'error': 'Please upload a valid PDF file'}), 400

    detect_tables = request.form.get('detect_tables', 'true') == 'true'

    uid = uuid.uuid4().hex
    pdf_path  = os.path.join(UPLOAD_FOLDER, f"{uid}_input.pdf")
    docx_path = os.path.join(UPLOAD_FOLDER, f"{uid}_output.docx")

    try:
        file.save(pdf_path)
        pages = convert_pdf_to_word(pdf_path, docx_path, detect_tables)
        original_name = os.path.splitext(file.filename)[0]
        return jsonify({
            'success': True,
            'pages': pages,
            'download_id': uid,
            'filename': f"{original_name}.docx"
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)


@app.route('/download/<uid>')
def download(uid):
    docx_path = os.path.join(UPLOAD_FOLDER, f"{uid}_output.docx")
    filename  = request.args.get('filename', 'converted.docx')
    if not os.path.exists(docx_path):
        return jsonify({'error': 'File not found or expired'}), 404
    return send_file(docx_path, as_attachment=True, download_name=filename)


if __name__ == '__main__':
    app.run(debug=True)
